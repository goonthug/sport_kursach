"""
Генератор договоров в формате ГОСТ Р 7.0.97-2016.

Шрифт Times New Roman 14pt, межстрочный 1.5, поля 30/15/20/20 мм,
абзацный отступ 1.25 см, текст по ширине. Нумерация страниц внизу по центру.
"""
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


_FONT = 'Times New Roman'
_SIZE = Pt(14)
_SIZE_SMALL = Pt(12)
_SPACING = 1.5
_INDENT = Cm(1.25)
_BLACK = RGBColor(0, 0, 0)

_MONTHS_RU = {
    'January': 'января', 'February': 'февраля', 'March': 'марта',
    'April': 'апреля', 'May': 'мая', 'June': 'июня',
    'July': 'июля', 'August': 'августа', 'September': 'сентября',
    'October': 'октября', 'November': 'ноября', 'December': 'декабря',
}


def _ru_date(d) -> str:
    s = d.strftime('%d %B %Y')
    for en, ru in _MONTHS_RU.items():
        s = s.replace(en, ru)
    return s


def _apply_font(run, size=None, bold=False):
    run.font.name = _FONT
    run.font.size = size or _SIZE
    run.font.bold = bold
    run.font.color.rgb = _BLACK


def _fmt(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         indent=_INDENT, before=0, after=0):
    pf = paragraph.paragraph_format
    pf.alignment = align
    pf.first_line_indent = indent
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = _SPACING


def _para(doc, text='', align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          bold=False, indent=_INDENT, before=0, after=0):
    p = doc.add_paragraph()
    _fmt(p, align, indent, before, after)
    if text:
        run = p.add_run(text)
        _apply_font(run, bold=bold)
    return p


def _section_header(doc, text, before=12):
    return _para(doc, text.upper(), align=WD_ALIGN_PARAGRAPH.CENTER,
                 bold=True, indent=Cm(0), before=before, after=6)


def _sub(doc, number, text):
    return _para(doc, f'{number}. {text}')


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _add_page_numbers(doc):
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _field(run, name):
        begin = OxmlElement('w:fldChar')
        begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(begin)
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = f' {name} '
        run._r.append(instr)
        end = OxmlElement('w:fldChar')
        end.set(qn('w:fldCharType'), 'end')
        run._r.append(end)

    for text, is_field, fname in [
        ('Стр. ', False, None),
        ('', True, 'PAGE'),
        (' из ', False, None),
        ('', True, 'NUMPAGES'),
    ]:
        r = fp.add_run(text)
        _apply_font(r, size=_SIZE_SMALL)
        if is_field:
            _field(r, fname)


def _city_date_row(doc, city, date_str):
    table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(table)
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    for p in (left, right):
        _fmt(p, indent=Cm(0), before=6, after=6)
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _apply_font(left.add_run(f'г. {city}'))
    _apply_font(right.add_run(f'{date_str} г.'))


def _setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)

    normal = doc.styles['Normal']
    normal.font.name = _FONT
    normal.font.size = _SIZE
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = _SPACING
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    _add_page_numbers(doc)
    return doc


# ПРАВКА 1: left_has_seal / right_has_seal — М.П. только у стороны с печатью
def _signature_table(doc, left_title, left_name, left_extra,
                     right_title, right_name, right_extra,
                     left_has_seal=True, right_has_seal=True):
    _para(doc, '', before=6)
    _section_header(doc, 'РЕКВИЗИТЫ И ПОДПИСИ СТОРОН', before=12)

    table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(table)

    def _fill(cell, title, name, extras, has_seal):
        first = cell.paragraphs[0]
        _fmt(first, align=WD_ALIGN_PARAGRAPH.LEFT, indent=Cm(0))
        r = first.add_run(title)
        _apply_font(r, bold=True)

        for line in [name] + extras:
            p2 = cell.add_paragraph()
            _fmt(p2, align=WD_ALIGN_PARAGRAPH.LEFT, indent=Cm(0))
            _apply_font(p2.add_run(line))

        blank = cell.add_paragraph()
        _fmt(blank, indent=Cm(0), before=12)

        sig = cell.add_paragraph()
        _fmt(sig, align=WD_ALIGN_PARAGRAPH.LEFT, indent=Cm(0))
        _apply_font(sig.add_run('Подпись: ___________________'))

        if has_seal:
            mp = cell.add_paragraph()
            _fmt(mp, align=WD_ALIGN_PARAGRAPH.LEFT, indent=Cm(0))
            _apply_font(mp.add_run('М.П.'))

    _fill(table.cell(0, 0), left_title, left_name, left_extra, left_has_seal)
    _fill(table.cell(0, 1), right_title, right_name, right_extra, right_has_seal)


# ---------------------------------------------------------------------------
# Публичный интерфейс
# ---------------------------------------------------------------------------

def generate_rental_contract(
    *,
    contract_number: str,
    contract_city: str,
    today,
    manager_full_name: str,
    client_full_name: str,
    passport_series: str,
    passport_number: str,
    passport_issue_date,
    passport_department_code: str,
    inventory_name: str,
    rental_days: int,
    start_date,
    end_date,
    price_per_day,
    total_price,
    deposit_amount,
) -> Document:
    """
    Договор аренды спортивного инвентаря (клиент <-> менеджер).
    Все параметры -- скалярные значения, не Django-объекты.
    """
    import re
    # ПРАВКА 3: убираем возможный префикс 'N ' или 'No ' если number пришёл из старой БД
    contract_number = re.sub(r'^[Nn№o\.]+\s*', '', str(contract_number)).strip()

    doc = _setup_doc()
    date_str = _ru_date(today)

    # Заголовок
    _para(doc, 'ДОГОВОР АРЕНДЫ СПОРТИВНОГО ИНВЕНТАРЯ',
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=Cm(0), after=0)
    # ПРАВКА 3: символ № (U+2116), не буква N
    _para(doc, f'№ {contract_number}',
          align=WD_ALIGN_PARAGRAPH.CENTER, indent=Cm(0), after=4)
    _city_date_row(doc, contract_city, date_str)

    # Преамбула — ПРАВКА 4: ёлочки вместо прямых кавычек
    passport_str = f'паспорт: серия {passport_series} № {passport_number}'
    if passport_issue_date:
        issued = (passport_issue_date.strftime('%d.%m.%Y')
                  if hasattr(passport_issue_date, 'strftime') else str(passport_issue_date))
        passport_str += f', выдан {issued}'
    if passport_department_code:
        passport_str += f', код подразделения {passport_department_code}'

    _para(doc,
          f'{manager_full_name}, именуемый в дальнейшем «Арендодатель», с одной стороны, '
          f'и {client_full_name}, {passport_str}, именуемый(ая) в дальнейшем «Арендатор», '
          f'с другой стороны, совместно именуемые «Стороны», заключили настоящий договор '
          f'о нижеследующем:',
          before=6, after=6)

    # 1. ПРЕДМЕТ ДОГОВОРА
    _section_header(doc, '1. ПРЕДМЕТ ДОГОВОРА')

    _sub(doc, '1.1',
         f'Арендодатель обязуется предоставить Арендатору во временное владение и пользование '
         f'спортивный инвентарь: {inventory_name} (далее — Инвентарь), а Арендатор обязуется '
         f'принять Инвентарь, уплатить арендную плату и возвратить его по истечении срока аренды.')

    start_s = start_date.strftime('%d.%m.%Y') if hasattr(start_date, 'strftime') else str(start_date)
    end_s = end_date.strftime('%d.%m.%Y') if hasattr(end_date, 'strftime') else str(end_date)
    _sub(doc, '1.2',
         f'Срок аренды составляет {rental_days} дней, с {start_s} по {end_s} включительно.')
    _sub(doc, '1.3',
         f'Место передачи Инвентаря: г. {contract_city}.')

    # 2. ПРАВА И ОБЯЗАННОСТИ СТОРОН
    _section_header(doc, '2. ПРАВА И ОБЯЗАННОСТИ СТОРОН')

    _sub(doc, '2.1', 'Арендодатель обязан:')
    _sub(doc, '2.1.1',
         'Передать Инвентарь Арендатору в надлежащем техническом состоянии в согласованный срок.')
    _sub(doc, '2.1.2',
         'Обеспечить Арендатора необходимыми инструкциями по использованию Инвентаря.')
    _sub(doc, '2.1.3',
         'Устранять за свой счёт недостатки Инвентаря, обнаруженные не по вине Арендатора.')

    _sub(doc, '2.2', 'Арендатор обязан:')
    _sub(doc, '2.2.1',
         'Использовать Инвентарь по назначению и в соответствии с его техническими характеристиками.')
    _sub(doc, '2.2.2',
         'Обеспечить сохранность Инвентаря и нести ответственность за его повреждение или утрату.')
    _sub(doc, '2.2.3',
         'Возвратить Инвентарь в установленный срок в надлежащем состоянии с учётом нормального износа.')
    _sub(doc, '2.2.4',
         'Своевременно вносить арендную плату в установленном настоящим договором порядке.')

    # 3. АРЕНДНАЯ ПЛАТА
    _section_header(doc, '3. АРЕНДНАЯ ПЛАТА И ПОРЯДОК РАСЧЕТОВ')

    _sub(doc, '3.1',
         f'Арендная плата за пользование Инвентарем составляет {price_per_day} рублей в сутки.')
    _sub(doc, '3.2',
         f'Общая сумма арендной платы за весь срок аренды составляет {total_price} рублей.')
    _sub(doc, '3.3',
         f'Залоговая сумма составляет {deposit_amount} рублей и вносится Арендатором '
         f'при подписании настоящего договора. Залог возвращается при возврате Инвентаря '
         f'в надлежащем состоянии.')
    _sub(doc, '3.4',
         'Оплата производится через электронную систему платформы «СпортРент».')

    # 4. ОТВЕТСТВЕННОСТЬ СТОРОН
    _section_header(doc, '4. ОТВЕТСТВЕННОСТЬ СТОРОН')

    _sub(doc, '4.1',
         'В случае повреждения или утраты Инвентаря Арендатор возмещает Арендодателю '
         'реальный ущерб в полном объёме.')
    # ПРАВКА 2: плата за продление вместо двойной неустойки
    _sub(doc, '4.2',
         'За каждый день просрочки возврата Инвентаря Арендатор '
         'обязан уплатить плату за продление аренды в размере '
         'суточной арендной платы, установленной п. 3.1 настоящего '
         'договора.')
    _sub(doc, '4.3',
         'Стороны освобождаются от ответственности за неисполнение обязательств, '
         'если такое неисполнение вызвано обстоятельствами непреодолимой силы (форс-мажор).')

    # 5. РАЗРЕШЕНИЕ СПОРОВ
    _section_header(doc, '5. РАЗРЕШЕНИЕ СПОРОВ')

    _sub(doc, '5.1',
         'Все споры и разногласия, которые могут возникнуть между Сторонами, '
         'разрешаются путём переговоров.')
    _sub(doc, '5.2',
         'При невозможности разрешения споров путём переговоров Стороны вправе обратиться '
         'в суд в соответствии с действующим законодательством Российской Федерации.')

    # 6. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ
    _section_header(doc, '6. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ')

    _sub(doc, '6.1',
         'Настоящий договор составлен в двух экземплярах, имеющих одинаковую юридическую '
         'силу, по одному для каждой из Сторон.')
    _sub(doc, '6.2',
         'Во всём, что не предусмотрено настоящим договором, Стороны руководствуются '
         'действующим законодательством Российской Федерации.')
    _sub(doc, '6.3',
         'Любые изменения и дополнения к настоящему договору действительны при условии, '
         'если они совершены в письменной форме и подписаны обеими Сторонами.')

    # Подписи (ПРАВКИ 1 + 4)
    passport_short = f'Паспорт: серия {passport_series} № {passport_number}'
    _signature_table(
        doc,
        left_title='Арендодатель:',
        left_name=manager_full_name,
        left_extra=[],
        right_title='Арендатор:',
        right_name=client_full_name,
        right_extra=[passport_short],
        left_has_seal=True,
        right_has_seal=False,
    )

    return doc


def generate_owner_contract(
    *,
    contract_city: str,
    today,
    owner_full_name: str,
    manager_full_name: str,
    inventory_name: str,
    bank_name: str = '',
    account_number: str = '',
    recipient_name: str = '',
) -> Document:
    """
    Агентский договор (владелец <-> платформа).
    Платформа выступает агентом от имени и за счёт Владельца.
    """
    doc = _setup_doc()
    date_str = _ru_date(today)

    # Заголовок
    _para(doc, 'АГЕНТСКИЙ ДОГОВОР',
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=Cm(0), after=4)
    _city_date_row(doc, contract_city, date_str)

    # Преамбула — ПРАВКА 4: «ёлочки»
    _para(doc,
          f'{owner_full_name}, именуемый в дальнейшем «Владелец», с одной стороны, '
          f'и {manager_full_name}, действующий от имени платформы «СпортРент» (далее — Платформа), '
          f'выступающей в роли агента, действующего от имени и за счёт Владельца, '
          f'именуемый в дальнейшем «Агент», с другой стороны, заключили настоящий договор '
          f'о нижеследующем:',
          before=6, after=6)

    # 1. ПРЕДМЕТ ДОГОВОРА
    _section_header(doc, '1. ПРЕДМЕТ ДОГОВОРА')

    _sub(doc, '1.1',
         f'Владелец поручает Агенту, а Агент принимает на себя обязательство от имени '
         f'и за счёт Владельца совершать действия по сдаче в аренду спортивного инвентаря: '
         f'{inventory_name} (далее — Инвентарь) на платформе «СпортРент» третьим лицам.')
    _sub(doc, '1.2',
         'Агент осуществляет поиск арендаторов, приём платежей, организацию передачи '
         'и возврата Инвентаря от имени Владельца.')

    # 2. ПРАВА И ОБЯЗАННОСТИ СТОРОН
    _section_header(doc, '2. ПРАВА И ОБЯЗАННОСТИ СТОРОН')

    _sub(doc, '2.1', 'Владелец обязан:')
    _sub(doc, '2.1.1',
         'Предоставить достоверные сведения об Инвентаре, включая его техническое '
         'состояние и характеристики.')
    _sub(doc, '2.1.2',
         'Своевременно передавать Инвентарь арендаторам в надлежащем состоянии.')
    _sub(doc, '2.1.3',
         'Уведомлять Агента о недоступности Инвентаря не менее чем за 48 часов.')

    _sub(doc, '2.2', 'Агент обязан:')
    _sub(doc, '2.2.1',
         'Осуществлять продвижение и размещение информации об Инвентаре на платформе.')
    _sub(doc, '2.2.2',
         'Производить расчёты с Владельцем в установленные настоящим договором сроки.')
    _sub(doc, '2.2.3',
         'Информировать Владельца о состоянии его Инвентаря и результатах аренд.')

    # 3. АГЕНТСКОЕ ВОЗНАГРАЖДЕНИЕ
    _section_header(doc, '3. АГЕНТСКОЕ ВОЗНАГРАЖДЕНИЕ И ПОРЯДОК РАСЧЕТОВ')

    _sub(doc, '3.1',
         'Агентское вознаграждение составляет 30 % от суммы каждой арендной сделки, '
         'совершённой Агентом от имени Владельца.')
    _sub(doc, '3.2',
         'Оставшиеся 70 % перечисляются Владельцу в течение 3 (трёх) рабочих дней после завершения аренды.')

    req_parts = []
    if recipient_name:
        req_parts.append(f'получатель: {recipient_name}')
    if bank_name:
        req_parts.append(f'банк: {bank_name}')
    if account_number:
        req_parts.append(f'номер счёта: {account_number}')
    req_text = (f'Банковские реквизиты Владельца: {"; ".join(req_parts)}.'
                if req_parts else 'Банковские реквизиты Владельца указываются в приложении.')
    _sub(doc, '3.3', req_text)

    # 4. ОТВЕТСТВЕННОСТЬ СТОРОН
    _section_header(doc, '4. ОТВЕТСТВЕННОСТЬ СТОРОН')

    _sub(doc, '4.1',
         'Агент не несёт ответственности за действия арендаторов, повлекшие повреждение '
         'или утрату Инвентаря, при условии соблюдения установленных процедур проверки.')
    _sub(doc, '4.2',
         'Владелец несёт ответственность за достоверность сведений об Инвентаре и его '
         'техническое состояние на момент передачи арендатору.')
    _sub(doc, '4.3',
         'Стороны освобождаются от ответственности за неисполнение обязательств, '
         'вызванное обстоятельствами непреодолимой силы (форс-мажор).')

    # 5. СРОК ДЕЙСТВИЯ ДОГОВОРА
    _section_header(doc, '5. СРОК ДЕЙСТВИЯ ДОГОВОРА')

    _sub(doc, '5.1',
         'Настоящий договор вступает в силу с момента его подписания Сторонами и '
         'действует до момента прекращения сотрудничества.')
    _sub(doc, '5.2',
         'Каждая из Сторон вправе в одностороннем порядке отказаться от исполнения '
         'настоящего договора, уведомив другую Сторону за 30 (тридцать) дней.')

    # 6. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ
    _section_header(doc, '6. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ')

    _sub(doc, '6.1',
         'Настоящий договор составлен в двух экземплярах, имеющих одинаковую юридическую '
         'силу, по одному для каждой из Сторон.')
    _sub(doc, '6.2',
         'Во всём, что не предусмотрено настоящим договором, Стороны руководствуются '
         'действующим законодательством Российской Федерации.')

    # Подписи (ПРАВКА 1)
    owner_extra = []
    if account_number:
        owner_extra.append(f'Счёт: {account_number}')
    if bank_name:
        owner_extra.append(f'Банк: {bank_name}')

    _signature_table(
        doc,
        left_title='Владелец:',
        left_name=owner_full_name,
        left_extra=owner_extra,
        right_title='Агент (Платформа):',
        right_name=manager_full_name,
        right_extra=[],
        left_has_seal=False,
        right_has_seal=True,
    )

    return doc
