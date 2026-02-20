"""
Views для кастомной административной панели.
Включает статистику, управление пользователями, модерацию.
"""

import logging
from datetime import datetime, timedelta
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.db.models import Count, Sum, Avg, Q
from django.core.paginator import Paginator
from django.utils import timezone

from users.models import User, Client, Owner, Manager
from inventory.models import Inventory, SportCategory
from rentals.models import Rental, Payment
from reviews.models import Review

logger = logging.getLogger('custom_admin')


def is_staff(user):
    """Проверка прав доступа к админ-панели."""
    return user.is_authenticated and user.role in ['manager', 'administrator']


@login_required
@user_passes_test(is_staff)
def admin_dashboard(request):
    """Главная страница административной панели с статистикой."""
    user = request.user
    manager_profile = None
    if user.role == 'manager' and hasattr(user, 'manager_profile'):
        manager_profile = user.manager_profile

    # Базовая статистика
    total_users = User.objects.count()
    total_clients = Client.objects.count()
    total_owners = Owner.objects.count()
    total_inventory = Inventory.objects.count()
    available_inventory = Inventory.objects.filter(status='available').count()
    pending_inventory = Inventory.objects.filter(status='pending').count()

    rentals_scope = Rental.objects.all()
    if manager_profile:
        rentals_scope = rentals_scope.filter(manager=manager_profile)

    total_rentals = rentals_scope.count()
    active_rentals = rentals_scope.filter(status='active').count()
    pending_rentals = rentals_scope.filter(status='pending').count()
    completed_rentals = rentals_scope.filter(status='completed').count()

    pending_reviews = Review.objects.filter(status='pending').count()

    # Финансовая статистика
    total_revenue = Payment.objects.filter(status='completed').aggregate(
        total=Sum('amount')
    )['total'] or 0

    # Статистика за последний месяц
    last_month = timezone.now() - timedelta(days=30)
    new_users_month = User.objects.filter(registration_date__gte=last_month).count()
    new_rentals_month = rentals_scope.filter(created_date__gte=last_month).count()

    # Популярные категории
    popular_categories = SportCategory.objects.annotate(
        total_items=Count('items'),
        total_rentals=Count('items__rentals')
    ).order_by('-total_rentals')[:5]

    # Топ инвентаря
    top_inventory = Inventory.objects.filter(
        status='available'
    ).annotate(
        rental_count=Count('rentals')
    ).order_by('-rental_count')[:5]

    # Недавняя активность
    recent_rentals = rentals_scope.select_related(
        'inventory', 'client', 'client__user'
    ).order_by('-created_date')[:10]

    context = {
        'total_users': total_users,
        'total_clients': total_clients,
        'total_owners': total_owners,
        'total_inventory': total_inventory,
        'available_inventory': available_inventory,
        'pending_inventory': pending_inventory,
        'total_rentals': total_rentals,
        'active_rentals': active_rentals,
        'pending_rentals': pending_rentals,
        'completed_rentals': completed_rentals,
        'pending_reviews': pending_reviews,
        'total_revenue': total_revenue,
        'new_users_month': new_users_month,
        'new_rentals_month': new_rentals_month,
        'popular_categories': popular_categories,
        'top_inventory': top_inventory,
        'recent_rentals': recent_rentals,
    }

    return render(request, 'custom_admin/dashboard.html', context)


@login_required
@user_passes_test(is_staff)
def admin_users(request):
    """Управление пользователями."""

    # Фильтрация
    role = request.GET.get('role')
    status = request.GET.get('status')
    search = request.GET.get('search', '').strip()

    users = User.objects.all().order_by('-registration_date')

    if role:
        users = users.filter(role=role)

    if status:
        users = users.filter(status=status)

    if search:
        users = users.filter(
            Q(email__icontains=search) |
            Q(phone__icontains=search)
        )

    # Пагинация
    paginator = Paginator(users, 20)
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)

    context = {
        'page_obj': page_obj,
        'selected_role': role,
        'selected_status': status,
        'search_query': search,
        'role_choices': User.ROLE_CHOICES,
        'status_choices': User.STATUS_CHOICES,
    }

    return render(request, 'custom_admin/users.html', context)


@login_required
@user_passes_test(is_staff)
def admin_user_block(request, user_id):
    """Блокировка пользователя."""

    if request.user.role != 'administrator':
        messages.error(request, 'Недостаточно прав')
        return redirect('custom_admin:users')

    user = get_object_or_404(User, user_id=user_id)

    if user.role == 'administrator':
        messages.error(request, 'Нельзя заблокировать администратора')
        return redirect('custom_admin:users')

    if request.method == 'POST':
        # Основная причина выбирается из списка, при необходимости админ может добавить комментарий.
        reason_choice = (request.POST.get('block_reason_choice') or '').strip()
        reason_custom = (request.POST.get('block_reason_custom') or '').strip()

        if not reason_choice and not reason_custom:
            messages.error(request, 'Укажите причину блокировки')
            return redirect('custom_admin:users')

        if reason_choice and reason_custom:
            reason = f'{reason_choice}. Дополнительно: {reason_custom}'
        elif reason_choice:
            reason = reason_choice
        else:
            reason = reason_custom

        user.status = 'blocked'
        user.block_reason = reason
        user.save()

        logger.info(f'Пользователь заблокирован: {user.email}, причина: {reason}, администратор {request.user.email}')
        messages.success(request, f'Пользователь {user.email} заблокирован')

    return redirect('custom_admin:users')


@login_required
@user_passes_test(is_staff)
def admin_user_unblock(request, user_id):
    """Разблокировка пользователя."""

    if request.user.role != 'administrator':
        messages.error(request, 'Недостаточно прав')
        return redirect('custom_admin:users')

    user = get_object_or_404(User, user_id=user_id)

    if request.method == 'POST':
        user.status = 'active'
        user.block_reason = ''
        user.save()

        logger.info(f'Пользователь разблокирован: {user.email} администратором {request.user.email}')
        messages.success(request, f'Пользователь {user.email} разблокирован')

    return redirect('custom_admin:users')


@login_required
@user_passes_test(is_staff)
def admin_inventory(request):
    """Управление инвентарем (для менеджеров)."""

    # Фильтрация
    status = request.GET.get('status')
    category = request.GET.get('category')
    search = request.GET.get('search', '').strip()

    inventory_qs = Inventory.objects.select_related('category', 'owner', 'manager').order_by('-added_date')

    if status:
        inventory_qs = inventory_qs.filter(status=status)

    if category:
        inventory_qs = inventory_qs.filter(category_id=category)

    if search:
        inventory_qs = inventory_qs.filter(
            Q(name__icontains=search) |
            Q(brand__icontains=search) |
            Q(model__icontains=search)
        )

    # Пагинация
    paginator = Paginator(inventory_qs, 20)
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)

    categories = SportCategory.objects.all()

    context = {
        'page_obj': page_obj,
        'selected_status': status,
        'selected_category': category,
        'search_query': search,
        'categories': categories,
    }

    return render(request, 'custom_admin/inventory.html', context)


@login_required
@user_passes_test(is_staff)
def admin_inventory_pending_detail(request, pk):
    """
    Просмотр заявки на инвентарь (полная карточка как в каталоге).
    Только для менеджера, для статусов pending и awaiting_contract.
    """
    if request.user.role != 'manager':
        messages.error(request, 'Недостаточно прав')
        return redirect('custom_admin:inventory')

    inventory = get_object_or_404(
        Inventory.objects.select_related('category', 'owner', 'owner__user', 'manager').prefetch_related('photos'),
        pk=pk
    )
    if inventory.status not in ['pending', 'awaiting_contract']:
        messages.warning(request, 'Эта заявка уже обработана')
        return redirect('custom_admin:inventory')

    # Проверяем, что менеджер работает с этим инвентарем (для статуса awaiting_contract)
    if inventory.status == 'awaiting_contract' and inventory.manager != request.user.manager_profile:
        messages.error(request, 'Вы не можете просматривать эту заявку')
        return redirect('custom_admin:inventory')

    return render(request, 'custom_admin/inventory_pending_detail.html', {'inventory': inventory})


@login_required
@user_passes_test(is_staff)
def admin_inventory_approve(request, pk):
    """Одобрение инвентаря менеджером с указанием залога."""

    if request.user.role != 'manager':
        messages.error(request, 'Недостаточно прав')
        return redirect('custom_admin:inventory')

    inventory = get_object_or_404(Inventory, pk=pk)

    if inventory.status != 'pending':
        messages.warning(request, 'Этот инвентарь уже обработан')
        return redirect('custom_admin:inventory')

    if request.method == 'POST':
        from decimal import Decimal
        deposit_amount = request.POST.get('deposit_amount', 0)
        try:
            deposit_amount = Decimal(str(deposit_amount)) if deposit_amount else Decimal('0')
            if deposit_amount < 0:
                deposit_amount = Decimal('0')
        except Exception:
            deposit_amount = Decimal('0')

        inventory.status = 'awaiting_contract'
        inventory.manager = request.user.manager_profile
        inventory.deposit_amount = deposit_amount
        inventory.save()

        logger.info(f'Инвентарь одобрен: {inventory.name}, залог {deposit_amount} менеджером {request.user.email}')
        messages.success(request, 'Инвентарь одобрен. Ожидается подписание договора.')

    return redirect('custom_admin:inventory')


@login_required
@user_passes_test(is_staff)
def admin_inventory_reject(request, pk):
    """Отклонение инвентаря менеджером."""

    if request.user.role != 'manager':
        messages.error(request, 'Недостаточно прав')
        return redirect('custom_admin:inventory')

    inventory = get_object_or_404(Inventory, pk=pk)

    if inventory.status != 'pending':
        messages.warning(request, 'Этот инвентарь уже обработан')
        return redirect('custom_admin:inventory')

    if request.method == 'POST':
        reason = request.POST.get('reason', '')
        inventory.status = 'rejected'
        inventory.rejection_reason = reason
        inventory.save()

        logger.info(f'Инвентарь отклонен: {inventory.name} причина: {reason}')
        messages.info(request, 'Инвентарь отклонен')

    return redirect('custom_admin:inventory')


@login_required
@user_passes_test(is_staff)
def admin_inventory_publish(request, pk):
    """Публикация инвентаря в каталоге (изменение статуса на 'available')."""

    if request.user.role != 'manager':
        messages.error(request, 'Недостаточно прав')
        return redirect('custom_admin:inventory')

    inventory = get_object_or_404(Inventory, pk=pk)

    if inventory.status != 'awaiting_contract':
        messages.warning(request, 'Инвентарь должен быть в статусе "Принят и ожидается подписание договора"')
        return redirect('custom_admin:inventory')

    if request.method == 'POST':
        # Проверяем, что менеджер работает с этим инвентарем
        if inventory.manager != request.user.manager_profile:
            messages.error(request, 'Вы не можете опубликовать этот инвентарь')
            return redirect('custom_admin:inventory')

        inventory.status = 'available'
        inventory.save()

        logger.info(f'Инвентарь опубликован в каталоге: {inventory.name} менеджером {request.user.email}')
        messages.success(request, 'Инвентарь опубликован в каталоге')

    return redirect('custom_admin:inventory')


@login_required
@user_passes_test(is_staff)
def inventory_contract_download(request, pk):
    """
    Скачивание договора аренды между владельцем инвентаря и менеджером (магазином).
    В договор подставляются данные из инвентаря и профилей владельца и менеджера.
    """
    inventory = get_object_or_404(
        Inventory.objects.select_related(
            'owner',
            'owner__user',
            'manager',
            'manager__user',
        ),
        pk=pk,
    )

    user = request.user

    # Доступ только менеджеру, который ведёт этот инвентарь, или администратору
    if user.role == 'manager':
        if not hasattr(user, 'manager_profile') or inventory.manager != user.manager_profile:
            messages.error(request, 'Недостаточно прав для скачивания договора')
            return redirect('custom_admin:inventory_pending_detail', pk=pk)
    elif user.role != 'administrator':
        messages.error(request, 'Недостаточно прав для скачивания договора')
        return redirect('custom_admin:inventory_pending_detail', pk=pk)

    if inventory.status != 'awaiting_contract':
        messages.error(request, 'Договор доступен только для инвентаря в статусе "Принят и ожидается подписание договора"')
        return redirect('custom_admin:inventory_pending_detail', pk=pk)

    # Определяем менеджера, который попадает в договор
    if user.role == 'manager' and hasattr(user, 'manager_profile'):
        manager_profile = user.manager_profile
    else:
        manager_profile = inventory.manager

    manager_full_name = manager_profile.full_name if manager_profile else ''

    owner = inventory.owner
    owner_full_name = owner.full_name

    # Получаем банковские реквизиты владельца
    bank_account = owner.bank_accounts.filter(is_default=True).first()
    if not bank_account:
        bank_account = owner.bank_accounts.first()

    bank_name = bank_account.bank_name if bank_account else ''
    account_number = bank_account.account_number if bank_account else ''
    recipient_name = bank_account.recipient_name if bank_account else ''

    today = timezone.now().date()

    # Путь к шаблону договора
    from pathlib import Path
    from django.conf import settings

    template_path = (Path(settings.BASE_DIR) / 'contracts' / 'dogovor_arendy_vladelec_magazin.docx')
    if not template_path.exists():
        messages.error(
            request,
            f'Файл шаблона договора не найден по пути: {template_path}. '
            f'Положите туда файл "Договор аренды между владельцем инвентаря и магазином(менеджером).docx".'
        )
        return redirect('custom_admin:inventory_pending_detail', pk=pk)

    try:
        from docx import Document
    except ImportError:
        messages.error(
            request,
            'Не установлен пакет python-docx. Установите его командой "pip install python-docx".'
        )
        return redirect('custom_admin:inventory_pending_detail', pk=pk)

    # Открываем шаблон и подставляем данные
    document = Document(str(template_path))

    def replace_in_paragraphs(find_text: str, replace_text: str):
        for paragraph in document.paragraphs:
            if find_text in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if find_text in inline[i].text:
                        inline[i].text = inline[i].text.replace(find_text, replace_text)

    # 1) Дата договора: ищем первую строку, где есть "г." (город) и подменяем на "г. Альметьевск, DD.MM.YYYY г."
    formatted_date = today.strftime('%d.%m.%Y')
    for paragraph in document.paragraphs:
        if 'г.' in paragraph.text:
            paragraph.text = f'г. Альметьевск, {formatted_date} г.'
            break

    # 2) ФИО владельца: после слов "Индивидуальный предприниматель"
    for paragraph in document.paragraphs:
        if 'Индивидуальный предприниматель' in paragraph.text:
            # Заменяем текст после "Индивидуальный предприниматель"
            text = paragraph.text
            if 'Индивидуальный предприниматель' in text:
                # Ищем место после "Индивидуальный предприниматель" и заменяем
                parts = text.split('Индивидуальный предприниматель', 1)
                if len(parts) > 1:
                    paragraph.text = f'Индивидуальный предприниматель {owner_full_name}{parts[1].lstrip()}'
                else:
                    paragraph.text = f'Индивидуальный предприниматель {owner_full_name}'
            break

    # 3) ФИО менеджера: после слов "с одной стороны, и"
    for paragraph in document.paragraphs:
        if 'с одной стороны, и' in paragraph.text.lower():
            # Оставляем фразу и добавляем ФИО менеджера
            text = paragraph.text
            if 'с одной стороны, и' in text.lower():
                # Находим позицию "с одной стороны, и" (регистронезависимо)
                idx = text.lower().find('с одной стороны, и')
                if idx != -1:
                    before = text[:idx + len('с одной стороны, и')]
                    after = text[idx + len('с одной стороны, и'):].lstrip()
                    paragraph.text = f'{before} {manager_full_name}{after}'
            break

    # 4) Адреса и реквизиты: ФИО владельца, номер счета, название банка, получатель
    for paragraph in document.paragraphs:
        if 'Адреса и реквизиты' in paragraph.text or 'Адреса и реквизиты:' in paragraph.text:
            # Ищем следующий абзац с реквизитами
            idx = document.paragraphs.index(paragraph)
            if idx + 1 < len(document.paragraphs):
                next_para = document.paragraphs[idx + 1]
                parts = [owner_full_name]
                if account_number:
                    parts.append(f'Номер счета: {account_number}')
                if bank_name:
                    parts.append(f'Название банка: {bank_name}')
                if recipient_name:
                    parts.append(f'Получатель: {recipient_name}')
                next_para.text = ', '.join(parts)
            break

    # Сохраняем договор в память и отдаём как .docx
    from io import BytesIO
    from django.http import HttpResponse

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="dogovor_arendy_{inventory.inventory_id}.docx"'

    return response


@login_required
@user_passes_test(is_staff)
def export_inventory_xlsx(request):
    """Экспорт инвентаря в XLSX."""
    try:
        from core.utils import export_inventory_to_xlsx
    except ImportError:
        messages.error(request, 'Модуль экспорта не найден. Создайте файл core/utils.py')
        return redirect('custom_admin:dashboard')

    inventory_qs = Inventory.objects.select_related('category', 'owner').all()

    # Применяем фильтры из GET параметров
    status = request.GET.get('status')
    if status:
        inventory_qs = inventory_qs.filter(status=status)

    return export_inventory_to_xlsx(inventory_qs)


@login_required
@user_passes_test(is_staff)
def export_inventory_pdf(request):
    """Экспорт инвентаря в PDF."""
    try:
        from core.utils import export_inventory_to_pdf
    except ImportError:
        messages.error(request, 'Модуль экспорта не найден. Создайте файл core/utils.py')
        return redirect('custom_admin:dashboard')

    inventory_qs = Inventory.objects.select_related('category').all()

    status = request.GET.get('status')
    if status:
        inventory_qs = inventory_qs.filter(status=status)

    return export_inventory_to_pdf(inventory_qs)


@login_required
@user_passes_test(is_staff)
def export_rentals_xlsx(request):
    """Экспорт аренд в XLSX."""
    try:
        from core.utils import export_rentals_to_xlsx
    except ImportError:
        messages.error(request, 'Модуль экспорта не найден. Создайте файл core/utils.py')
        return redirect('custom_admin:dashboard')

    rentals_qs = Rental.objects.select_related('inventory', 'client').all()

    status = request.GET.get('status')
    if status:
        rentals_qs = rentals_qs.filter(status=status)

    return export_rentals_to_xlsx(rentals_qs)


@login_required
@user_passes_test(is_staff)
def export_stats_pdf(request):
    """Экспорт статистики в PDF."""
    try:
        from core.utils import export_stats_to_pdf
    except ImportError:
        messages.error(request, 'Модуль экспорта не найден. Создайте файл core/utils.py')
        return redirect('custom_admin:dashboard')

    stats_data = {
        'total_users': User.objects.count(),
        'total_clients': Client.objects.count(),
        'total_owners': Owner.objects.count(),
        'total_inventory': Inventory.objects.count(),
        'available_inventory': Inventory.objects.filter(status='available').count(),
        'total_rentals': Rental.objects.count(),
        'active_rentals': Rental.objects.filter(status='active').count(),
        'completed_rentals': Rental.objects.filter(status='completed').count(),
        'total_revenue': Payment.objects.filter(status='completed').aggregate(
            total=Sum('amount')
        )['total'] or 0,
    }

    return export_stats_to_pdf(stats_data)